# -*- coding: utf-8 -*-
"""app.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1p7Ef13LULpVGcvXTjI7MQHdrFr51zMNb
"""

# === STREAMLIT DEPLOYMENT VERSION ===

# STEP 1: Import Required Libraries
import os
import re
import io
import json
import glob
import plotly.graph_objects as go
import pandas as pd
from collections import Counter
from collections import defaultdict
from PIL import Image
from pathlib import Path
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_unstructured import UnstructuredLoader
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.schema import Document as LCDocument
from PyPDF2 import PdfReader
from docx import Document
import warnings
import shutil
import streamlit as st
import streamlit.components.v1 as components


warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=pd.errors.ParserWarning)

# STEP 2: Load Environment Variables
load_dotenv()
IFI_API_KEY = os.getenv("IFI_API_KEY")  # <-- ADD YOUR API KEY to .streamlit/secrets.toml or env vars

# STEP 3: Ensure necessary folders exist before file operations
for folder in ['historical_documents', 'risks_document', 'target_document', 'outputs']:
    Path(folder).mkdir(parents=True, exist_ok=True)

# STEP 4: Define the RAG Risk Analysis Class
class RAGProcurementRisksAnalysis:
    def __init__(self, api_key, query, historical_documents_folder_path, risks_document_folder_path, target_document_folder_path, risk_analysis_output_path):
        self.api_key = api_key
        self.query = query
        self.historical_documents = self.load_documents(historical_documents_folder_path)
        self.risks_document = self.load_documents(risks_document_folder_path)
        self.target_document = self.load_documents(target_document_folder_path)
        self.risk_analysis_output_path = risk_analysis_output_path
    def extract_json_string(text):
        json_start = text.find('{')
        json_text = text[json_start:].strip()
        if json_text.endswith("```"):
            json_text = json_text[:-3].strip()
        return json_text

    def load_documents(self, folder_path):
        all_documents = []
        supported_exts = ["csv", "pdf", "docx"]
        for ext in supported_exts:
            files = glob.glob(f"{folder_path}/*.{ext}")
            for file_path in files:
                try:
                    if file_path.endswith(".csv"):
                        try:
                            with open(file_path, "r", encoding="utf-8") as f:
                                content = f.read()
                        except UnicodeDecodeError:
                            with open(file_path, "r", encoding="latin1") as f:
                                content = f.read()
                        doc = LCDocument(page_content=content)
                        all_documents.append(doc)
                    else:
                        loader = UnstructuredLoader(file_path=file_path)
                        documents = loader.load()
                        all_documents.extend(documents)
                except Exception as e:
                    print(f"⚠️ Could not load {file_path}: {e}")
        return all_documents

    def create_embeddings(self):
        embeddings = OpenAIEmbeddings(openai_api_key=self.api_key)
        vector_store = FAISS.from_documents(self.historical_documents, embeddings)
        return vector_store, embeddings

    def semantic_search(self):
        vector_store, embeddings = self.create_embeddings()
        query_embedding = embeddings.embed_query(self.query)
        risks_document_embedding = embeddings.embed_query(self.risks_document[0].page_content)
        target_document_embedding = embeddings.embed_query(self.target_document[0].page_content)
    
        retrieved_by_query = vector_store.similarity_search_by_vector(query_embedding, k=3)
        retrieved_by_risks = vector_store.similarity_search_by_vector(risks_document_embedding, k=3)
        retrieved_by_target = vector_store.similarity_search_by_vector(target_document_embedding, k=3)
    
        retrieved_documents = list({doc.page_content: doc for doc in retrieved_by_query + retrieved_by_target + retrieved_by_risks}.values())

        if not retrieved_documents:
            print("⚠️ No documents retrieved during semantic search!")
    
        print(f"🔍 Retrieved {len(retrieved_documents)} relevant docs for semantic search.")
    
        return "\n\n".join([f"Document {i + 1}: {doc.page_content}" for i, doc in enumerate(retrieved_documents)])

    def save_risk_analysis_to_file(self, risk_analysis):
        file_path = f"{self.risk_analysis_output_path}/risk_analysis.txt"
        os.makedirs(self.risk_analysis_output_path, exist_ok=True)
        with open(file_path, "w") as file:
            file.write(risk_analysis)
    
    def generate_risks_analysis_rag(self):
        llm = ChatOpenAI(model="gpt-4o", temperature=0.5, openai_api_key=self.api_key)
    
        # Validate document availability
        if not self.risks_document or not self.risks_document[0].page_content.strip():
            st.error("❌ The risks document is missing or empty.")
            return "Error: Risks document missing."
        
        if not self.target_document or not self.target_document[0].page_content.strip():
            st.error("❌ The target document is missing or empty.")
            return "Error: Target document missing."
    
        # ✅ Add this
        risks_content = self.risks_document[0].page_content
        target_content = self.target_document[0].page_content
    
        # Try semantic search and fallback
        retrieved_docs_str = self.semantic_search()
        if not retrieved_docs_str.strip():
            retrieved_docs_str = "No relevant documents were retrieved. Please proceed with only risks and target documents."
    
        if not risks_content.strip():
            st.error("❌ The risks document is empty. Please upload a valid file.")
            return "Error: Risks document is empty."
    
        if not target_content.strip():
            st.error("❌ The target document is empty. Please upload a valid file.")
            return "Error: Target document is empty."
    
        # Debug logs
        print("----- Prompt Preview -----")
        print("Query:", self.query)
        print("--- Retrieved Docs ---")
        print(retrieved_docs_str[:500])
        print("--- Risks Document ---")
        print(risks_content[:500])
        print("--- Target Document ---")
        print(target_content[:500])

    
        prompt_template = PromptTemplate(
            input_variables=["retrieved_docs_str", "risks_document_content", "target_document_content"],
            template="""
        You are an AI risk analyst reviewing procurement documentation.
        
        Analyze the risks based on the documents below.
        
        Retrieved Content:
        {retrieved_docs_str}
        
        Risks Document:
        {risks_document_content}
        
        Target Procurement Document:
        {target_document_content}
        
        Respond ONLY in JSON using this structure:
        {{
          "summary": {{
            "high": number of risks with "severity": "High",
            "medium": number of risks with "severity": "Medium",
            "low": number of risks with "severity": "Low",
            "budget_variance": "string",
            "schedule_variance": "string",
            "risk_score": "int"
          }},
          "risks": [
            {{
              "type": "string",
              "title": "string",
              "severity": "High" | "Medium" | "Low",
              "confidence": "int",
              "key_data": "string",
              "mitigation": "string"
            }}
          ],
          "timeline": [
            {{
              "task": "string",
              "planned_start": "YYYY-MM-DD",
              "planned_end": "YYYY-MM-DD",
              "actual_start": "YYYY-MM-DD",
              "actual_end": "YYYY-MM-DD",
              "risk": "string"
            }}
          ]
        }}
        
        Make sure the response is **valid JSON only**. Do not include any explanatory text or markdown.
        """
        )

        

    
        chain = LLMChain(llm=llm, prompt=prompt_template)
        print("📦 Inputs to chain.run():")
        print("retrieved_docs_str:", bool(retrieved_docs_str))
        print("risks_document_content:", bool(risks_content))
        print("target_document_content:", bool(target_content))

        inputs = {
            "retrieved_docs_str": retrieved_docs_str.strip(),
            "risks_document_content": risks_content.strip(),
            "target_document_content": target_content.strip()
        }
        
        # 🧪 Debugging: Print input keys and their lengths
        for k, v in inputs.items():
            print(f"[DEBUG] {k} - Length: {len(v)}")
        
        # Validate: Raise clear error if empty
        missing_keys = [k for k, v in inputs.items() if not v]
        if missing_keys:
            st.error(f"❌ Chain input error: Missing or empty keys: {missing_keys}")
            return f"Error: Chain input validation failed. Missing keys: {missing_keys}"
        
        
        try:
            response_obj = chain.invoke(inputs)
            # 1. Get LLM raw text response
            response_text = response_obj.get("text", "") if isinstance(response_obj, dict) else str(response_obj)
            
            # 2. Clean response
            response_text = response_text.strip()
            response_text = re.sub(r"^```(?:json)?", "", response_text, flags=re.MULTILINE)
            response_text = re.sub(r"```$", "", response_text, flags=re.MULTILINE)
            
            # 3. Strip leading non-JSON text
            json_start = response_text.find("{")
            if json_start == -1:
                st.error("❌ No JSON object detected in model output.")
                return response_text
            
            response_text = response_text[json_start:].strip()
            
            # 4. Parse JSON
            try:
                result_json = json.loads(response_text)
            except json.JSONDecodeError as e:
                st.warning("⚠️ Model response was not valid JSON. Showing raw response instead.")
                print("⚠️ JSONDecodeError:", e)
                print("🧾 Raw cleaned output:\n", response_text[:1000])
                st.session_state["raw_response_text"] = response_text
                return {
                    "summary": {},
                    "risks": [],
                    "timeline": []
                }, response_text

            
            # 5. Validate required keys
            required_keys = {"risks", "summary", "timeline"}
            missing_keys = required_keys - result_json.keys()
            if missing_keys:
                st.error(f"❌ JSON missing required keys: {missing_keys}")
                print("⚠️ JSON keys present:", result_json.keys())
                return json.dumps(result_json, indent=2)  # Fallback to raw view
            
            # 6. Save & return
            self.save_risk_analysis_to_file(json.dumps(result_json, indent=2))
            st.session_state["raw_response_text"] = response_text
            return result_json, response_text

            
        except Exception as e:
            st.error("❌ An unexpected error occurred while generating the risk analysis.")
            print("⚠️ Exception during chain.invoke:", e)
            return f"Error: {str(e)}"

# STEP 5: Preview Function
def preview_file(file, file_type, name="Uploaded file"):
    st.subheader(f"Preview: {name}")
    if file_type == "csv":
        df = pd.read_csv(file)
        st.dataframe(df.head())
    elif file_type == "pdf":
        reader = PdfReader(file)
        text = "\n".join([page.extract_text() for page in reader.pages[:2] if page.extract_text()])
        st.markdown("#### 📑 Extracted Preview with Highlights")
        highlighted_text = text[:2000].replace("risk", "**:red[risk]**").replace("delay", "**:orange[delay]**")
        st.markdown(highlighted_text, unsafe_allow_html=True)
    elif file_type == "docx":
        doc = Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
        st.text_area("DOCX Preview", text[:2000], height=200)

# STEP 6: Streamlit UI Setup
EXAMPLES_PATH = Path(__file__).resolve().parent.parent / "example_files"
st.set_page_config(page_title="MAESTRO", layout="centered")
# Load and display logo
logo_path = EXAMPLES_PATH / "maestro.png"
col_logo, col_title = st.columns([1, 4])  # Adjust width ratio if needed

with col_logo:
    if logo_path.exists():
        logo = Image.open(logo_path)
        st.image(logo, width=100)  # Optional: reduce width for balance

with col_title:
    st.markdown("<h1 style='margin-bottom: 0;'>MAESTRO</h1>", unsafe_allow_html=True)



st.sidebar.title("ℹ️ About")
st.sidebar.info('''
This tool uses LLM-based Retrieval-Augmented Generation (RAG) to assess risks in procurement documents.

- Built with LangChain + Streamlit
- Supports CSV, PDF, and DOCX
- Preview documents before analysis
- Securely runs in your environment
''')

st.markdown("### 📂 Upload Your Documents")
st.markdown("Drag & drop your files below. Supported: 📄 CSV, 📑 PDF, 📝 DOCX")

col1, col2, col3 = st.columns(3)
with col1:
    EXAMPLES_PATH = Path(__file__).resolve().parent.parent / "example_files"
    with open(EXAMPLES_PATH / "dataset1.csv", "rb") as f:
        st.download_button("📄 Sample History  Doc", f, file_name="dataset1.csv", help="Historical doc example")
with col2:
    with open(EXAMPLES_PATH / "risks.csv", "rb") as f:
        st.download_button("📑 Sample Risk Register", f, file_name="risks.csv", help="Risk types to reference")
with col3:
    with open(EXAMPLES_PATH / "dataset_no_risks.csv", "rb") as f:
        st.download_button("📝 Sample Target File", f, file_name="dataset_no_risks.csv", help="Target doc example")
    doc_labels = {
        "History Document": [],
        "Risk Register": [],
        "Target Procurement File": [],
    }

query = st.text_input("What do you want to know?", "What are the risks associated with this procurement document?")
    
uploaded_docs = {}

for label in doc_labels:
    st.markdown(f"**{label}:**")
    uploaded_files = st.file_uploader(
        f"Upload your {label}",
        type=["csv", "pdf", "docx"],
        key=label,
        accept_multiple_files=True,
        help={
            "History Document": "📚 Upload past procurement records. These help the model understand project patterns.",
            "Risk Register": "⚠️ Upload a file listing types of risks and descriptions (e.g., Risk Doc.csv).",
            "Target Procurement File": "🎯 Upload the project document you want to analyze (e.g., Target.csv)."
        }.get(label, "")
    )

    if uploaded_files:
        uploaded_docs[label] = []  # Initialize list to store files for this label

        for uploaded_file in uploaded_files:
            st.success(f"✅ Uploaded: {uploaded_file.name}")
            file_ext = uploaded_file.name.split(".")[-1]
            bytes_data = uploaded_file.getvalue()

            st.text(f"🧪 Uploaded {label}: {uploaded_file.name}, size: {len(bytes_data)} bytes")
            preview_file(io.BytesIO(bytes_data), file_ext, name=uploaded_file.name)

            uploaded_docs[label].append((uploaded_file.name, bytes_data))

# Extract specific categories if needed later
historical_files = uploaded_docs.get("History Document", [])
risks_files = uploaded_docs.get("Risk Register", [])
target_files = uploaded_docs.get("Target Procurement File", [])

# === Run Risk Analysis and Store in Session State ===
if st.button("Run Analysis"):
    if not IFI_API_KEY:
        st.error("Missing API key!")
    elif not historical_files or not risks_files or not target_files:
        st.warning("Please upload all required files.")
    else:
        with st.spinner("Processing files and analyzing..."):
            base_dir = Path(".")
            (base_dir / "historical_documents").mkdir(exist_ok=True)
            (base_dir / "risks_document").mkdir(exist_ok=True)
            (base_dir / "target_document").mkdir(exist_ok=True)

            for folder_name, file_list in [
                ("historical_documents", historical_files),
                ("risks_document", risks_files),
                ("target_document", target_files),
            ]:
                for fname, fbytes in file_list:
                    with open(base_dir / folder_name / fname, "wb") as out:
                        out.write(fbytes)

            rag = RAGProcurementRisksAnalysis(
                api_key=IFI_API_KEY,
                query=query,
                historical_documents_folder_path=base_dir / "historical_documents",
                risks_document_folder_path=base_dir / "risks_document",
                target_document_folder_path=base_dir / "target_document",
                risk_analysis_output_path=base_dir / "outputs"
            )

            st.session_state["risk_result"] = rag.generate_risks_analysis_rag()

def extract_risk_summary(text):
    summary = {
        "high": None,
        "medium": None,
        "low": None,
        "budget_variance": None,
        "schedule_variance": None,
        "risk_score": None
    }

    patterns = {
        "high": r"High Risks:\s*(\d+)",
        "medium": r"Medium Risks:\s*(\d+)",
        "low": r"Low Risks:\s*(\d+)",
        "budget_variance": r"Budget Variance:\s*([\$\d,]+(?: Overrun| Underrun)?)",
        "schedule_variance": r"Schedule Variance:\s*([^\n]+)",
        "risk_score": r"Risk Score:\s*(\d+)/100"
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text)
        if match:
            summary[key] = match.group(1)

    return summary

# === Render Analysis Results If Present ===
if "risk_result" in st.session_state:
    result_data, raw_output = st.session_state.get("risk_result", ({}, ""))

    if not isinstance(result_data, dict) or "risks" not in result_data or not isinstance(raw_output, str):
        st.error("⚠️ The model did not return a structured JSON output. Please try again or check the LLM output formatting.")
        st.markdown("### 🔍 Raw Output")
        st.markdown("### 📤 Raw LLM Output Before Parsing")
        st.code(result_data[:1000] if isinstance(result_data, str) else json.dumps(result_data, indent=2)[:1000])
        st.code(result_data if isinstance(result_data, str) else str(result_data))
    else:
        summary = result_data.get("summary", {})
        risks = result_data.get("risks", [])
        # ✅ Risk Score Calculation (confidence floor + weighted severity)
        # Weighted severity and confidence
        weights = {"high": 10, "medium": 5, "low": 1}
        total_score = 0
        
        for risk in risks:
            severity = risk["severity"].lower()
            confidence = risk.get("confidence", 100)
            weight = weights.get(severity, 0)
            total_score += weight * (confidence / 100)
        
        # Fixed max scale (e.g., assuming 10 high risks would hit 100)
        max_possible_score = 100
        risk_score_calc = min(int(total_score), max_possible_score)  # Cap at 100
        summary["risk_score"] = risk_score_calc
        
        # Optional debug print
        print(f"🧠 Final Risk Score: {risk_score_calc} from total: {total_score:.2f} / max: {max_score}")




        # Group and count risks

    grouped_risks = defaultdict(list)
    for risk in risks:
        grouped_risks[risk['severity'].lower()].append(risk)
    risk_counts = Counter(risk['severity'].lower() for risk in risks)
    
    # 📊 Risk Summary Panel
    with st.container():
        st.markdown("## 📊 Risk Summary")
        st.markdown("Quick overview of the identified risks and key project metrics.")
    
        summary_cols = st.columns([1, 1, 1])
        summary_cols[0].metric("🟥 High Risks", risk_counts.get("high", 0))
        summary_cols[1].metric("🟧 Medium Risks", risk_counts.get("medium", 0))
        summary_cols[2].metric("🟩 Low Risks", risk_counts.get("low", 0))
      
    # === 📋 Risk Explorer Tabs ===
    st.markdown("#### 📋 Risk Explorer")
    tabs = st.tabs([
        f"🟥 High Risks ({len(grouped_risks['high'])})",
        f"🟧 Medium Risks ({len(grouped_risks['medium'])})",
        f"🟩 Low Risks ({len(grouped_risks['low'])})"
    ])
    for i, severity in enumerate(["high", "medium", "low"]):
        with tabs[i]:
            if not grouped_risks[severity]:
                st.info("No risks found in this category.")
            for risk in grouped_risks[severity]:
                with st.expander(f"{risk['type']} — {severity.capitalize()} Risk ({risk['confidence']}%)", expanded=False):
                    st.markdown(f"**Key Insight:** {risk['key_data']}")
                    st.markdown(f"**Mitigation Plan:** {risk['mitigation']}")
                    st.markdown(
                        f"""<div title="Why this category?">{risk['severity']} based on: <b>{risk['key_data']}</b></div>""",
                        unsafe_allow_html=True
                    )

    st.markdown("---")
      
    st.markdown("### 📈 Variance Summary")
    with st.expander("ℹ️ How are these metrics calculated?", expanded=False):
        st.markdown("""
        - **📘 Budget Variance** is calculated by comparing the projected and actual costs found in the uploaded documents.
        - **⏱️ Schedule Variance** is based on delays between planned vs. actual milestones.
        - **🎯 Risk Score** is not a grade—**a higher number = greater risk**. It’s a weighted score from the model, reflecting risk count, severity, and confidence levels.
        """)
    col_left, col_right = st.columns([2, 3])
    
    with col_left:
        st.markdown(f"**📘 Budget Variance:**")
        st.success(f"{summary.get('budget_variance', 'N/A')}")
        st.markdown(f"**⏱️ Schedule Variance:**")
        st.success(f"{summary.get('schedule_variance', 'N/A')}")
    
    with col_right:
        score = summary.get("risk_score", 0)
        
        # Risk Score Title + Progress Bar
        st.markdown(f"🎯 **Overall Risk Level:** {score}/100")
        st.progress(score / 100)
    
        # Add vertical spacing before badge
        st.markdown("<div style='margin-top: 0.5rem'></div>", unsafe_allow_html=True)
    
        # Risk Level Label with Padding
        if score >= 75:
            st.markdown('<div style="background-color:#ffecec; padding: 0.75rem 1rem; border-radius: 10px; font-size: 1rem;"><span style="color: red; font-weight: bold;">🔴 High Risk Level</span></div>', unsafe_allow_html=True)
        elif score >= 40:
            st.markdown('<div style="background-color:#fff9e6; padding: 0.75rem 1rem; border-radius: 10px; font-size: 1rem;"><span style="color: #e69500; font-weight: bold;">🟠 Moderate Risk Level</span></div>', unsafe_allow_html=True)
        else:
            st.markdown('<div style="background-color:#e6f9ec; padding: 0.75rem 1rem; border-radius: 10px; font-size: 1rem;"><span style="color: #2e8b57; font-weight: bold;">🟢 Low Risk Level</span></div>', unsafe_allow_html=True)
    
        # Add space before expander
        st.markdown("<div style='margin-top: 0.5rem'></div>", unsafe_allow_html=True)
    
        # Risk Score Breakdown
        with st.expander("🧠 Risk Score Calculation Breakdown"):
            st.markdown(f"- **Total Weighted Score:** `{total_score:.2f}`")
            st.markdown(f"- **Max Possible Score:** `{max_score}`")
            st.markdown(f"- **Final Risk Score:** `{risk_score_calc}` out of 100")

    
    st.markdown("---")
    
    # === ⏱️ Timeline Section ===
    timeline_data = pd.DataFrame(result_data.get("timeline", []))
    if not timeline_data.empty:
        st.markdown("## ⏱️ Project Timeline")
        with st.expander("📅 View Timeline Chart", expanded=False):
            fig = go.Figure()
            
            # Bar: Planned Duration (baseline)
            fig.add_trace(go.Bar(
                y=timeline_data["task"],
                x=pd.to_datetime(timeline_data["planned_end"]) - pd.to_datetime(timeline_data["planned_start"]),
                base=timeline_data["planned_start"],
                orientation='h',
                name='Planned',
                marker_color='lightgray',
                hoverinfo='x+y'
            ))
            
            # Bar: Actual Duration (from target doc)
            fig.add_trace(go.Bar(
                y=timeline_data["task"],
                x=pd.to_datetime(timeline_data["actual_end"]) - pd.to_datetime(timeline_data["actual_start"]),
                base=timeline_data["actual_start"],
                orientation='h',
                name='Actual',
                marker_color='steelblue',
                hoverinfo='x+y'
            ))
            
            fig.update_layout(
                title="📅 Project Timeline: Planned vs. Actual",
                barmode='overlay',
                xaxis_title="Date",
                yaxis_title="Task",
                showlegend=True,
                height=400
            )
            
            st.plotly_chart(fig, use_container_width=True)



    st.markdown("### 📤 Export & Share")
    if isinstance(result_data, dict):
        export_text = json.dumps(result_data, indent=2)
    else:
        export_text = str(result_data)
    
    st.download_button("📄 Download as TXT", export_text, file_name="risk_analysis.txt")
    st.download_button("💾 Export as JSON", export_text, file_name="risk_analysis.json")
    st.session_state["jump_to"] = None
    

    
    # Fallback UI to debug raw output            
    if not isinstance(result_data, dict):
        st.error("⚠️ The model did not return a valid JSON. Check the prompt or document inputs.")
        st.markdown("### 🧾 Raw Output")
        st.code(result_data if isinstance(result_data, str) else str(result_data))
